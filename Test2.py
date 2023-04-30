from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


document = Document("abcd.docx")

paragraph = document.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = paragraph.add_run()
run.add_picture('upi.png', width=Inches(0.3), height=Inches(0.3), )

document.save('my_doc.docx')