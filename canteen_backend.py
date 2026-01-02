import os
import datetime
import qrcode
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ItemManager:
    def __init__(self, items_dir="Items"):
        self.items_dir = items_dir
        if not os.path.exists(self.items_dir):
            os.makedirs(self.items_dir)

    def get_item(self, code):
        """Retrieves item details by code. Returns dict or None."""
        file_path = os.path.join(self.items_dir, f"{code}.txt")
        if not os.path.exists(file_path):
            return None
        
        try:
            with open(file_path, "r") as f:
                lines = f.readlines()
                if len(lines) < 3:
                    return None
                return {
                    "code": code,
                    "name": lines[0].strip(),
                    "price": int(lines[1].strip()),
                    "stock": int(lines[2].strip())
                }
        except Exception as e:
            print(f"Error reading item {code}: {e}")
            return None

    def get_all_items(self):
        """Returns a list of all items."""
        items = []
        if not os.path.exists(self.items_dir):
            return items
            
        for filename in os.listdir(self.items_dir):
            if filename.endswith(".txt"):
                code = filename[:-4]
                item = self.get_item(code)
                if item:
                    items.append(item)
        return items

    def update_stock(self, code, quantity_sold):
        """Reduces stock by quantity_sold."""
        item = self.get_item(code)
        if not item:
            return False
        
        new_stock = item['stock'] - quantity_sold
        return self.save_item(code, item['name'], item['price'], new_stock)

    def save_item(self, code, name, price, stock):
        """Creates or updates an item."""
        file_path = os.path.join(self.items_dir, f"{code}.txt")
        try:
            with open(file_path, "w") as f:
                f.write(f"{name}\n{price}\n{stock}")
            return True
        except Exception as e:
            print(f"Error saving item {code}: {e}")
            return False

class InvoiceGenerator:
    def __init__(self, templates_dir="Templetes", invoices_dir="Invoices"):
        self.templates_dir = templates_dir
        self.invoices_dir = invoices_dir
        # Ensure directories exist
        if not os.path.exists(self.invoices_dir):
            os.makedirs(self.invoices_dir)

    def _get_next_invoice_number(self):
        count = 0
        for root, dirs, files in os.walk(self.invoices_dir):
            for file in files:
                if file.endswith('.docx'):
                    count += 1
        return count + 1

    def _replace_text(self, doc, old_text, new_text):
        for p in doc.paragraphs:
            if old_text in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        text = inline[i].text.replace(old_text, str(new_text))
                        inline[i].text = text

    def generate_invoice(self, items, discount, pay_mode):
        """
        items: list of dicts {'code', 'name', 'price', 'qty', 'total'}
        pay_mode: 'Cash' or 'UPI'
        """
        inv_no = self._get_next_invoice_number()
        now = datetime.datetime.now()
        current_date = now.strftime("%d-%m-%Y")
        current_time = now.strftime("%d-%m-%Y        Time - %H:%M:%S")
        
        # Create date folder
        date_folder = os.path.join(self.invoices_dir, current_date)
        if not os.path.exists(date_folder):
            os.makedirs(date_folder)

        # Select Template based on item count and payment mode
        # Logic adapted from original main.py:
        # 1 item -> 1.docx / 1QR.docx
        # 2 items -> 2.docx / 2QR.docx
        # etc.
        item_count = len(items)
        if item_count > 8:
            raise ValueError("Maximum 8 items allowed per invoice.")
            
        template_name = str(item_count)
        if pay_mode == 'UPI':
            template_name += "QR"
        if discount > 0:
             template_name += " D" # Assuming ' D' suffix for discount templates based on main.py
        
        template_path = os.path.join(self.templates_dir, f"{template_name}.docx")
        
        if not os.path.exists(template_path):
             # Fallback or error if specific template doesn't exist
             # For now, let's try to be robust or fail gracefully
             print(f"Warning: Template {template_path} not found.")
             # In a real app we might want a generic template, but here we rely on the existing ones.
             # If the user has exact filenames like '1.docx', '1QR.docx', '1 D.docx' etc.
        
        try:
            doc = Document(template_path)
        except Exception:
            raise FileNotFoundError(f"Template not found: {template_path}")

        # Replacements
        self._replace_text(doc, "111", f"INV {inv_no}")
        self._replace_text(doc, "112", current_time)
        
        grand_total = sum(item['total'] for item in items)
        final_amount = grand_total - discount
        total_qty = sum(item['qty'] for item in items)

        # Item details replacements
        # The original code used specific codes for specific lines (114, 115, 116 for item 1, etc.)
        # We need to map these dynamically.
        # Based on main.py:
        # Item 1: Name=114, Qty=115, Amount=116
        # Item 2: Name=117, Qty=118, Amount=119
        # Item 3: Name=120, Qty=121, Amount=122
        # ... and so on.
        
        base_codes = [114, 117, 120, 123, 126, 129, 132, 135] # Inferred pattern +3
        
        for i, item in enumerate(items):
            if i < len(base_codes):
                base = base_codes[i]
                self._replace_text(doc, str(base), item['name'])
                self._replace_text(doc, str(base+1), str(item['qty']))
                self._replace_text(doc, str(base+2), str(item['total']))

        self._replace_text(doc, "666", str(discount))
        self._replace_text(doc, "55", str(total_qty))
        self._replace_text(doc, "99", str(final_amount))

        # QR Code
        if pay_mode == 'UPI':
            upi_link = f"upi://pay?pa=nanigopalkaran77777@oksbi&pn=Tasty Confectionary&am={final_amount}&tn=Tasty Confectionary&cu=INR"
            qr_img = qrcode.make(upi_link)
            qr_path = "temp_qrcode.png"
            qr_img.save(qr_path)
            
            # Add QR to document
            # Original code adds it as a new paragraph at the end
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(qr_path, width=Inches(0.8), height=Inches(0.8))
            
        output_path = os.path.join(date_folder, f"{inv_no}.docx")
        doc.save(output_path)
        
        if pay_mode == 'UPI' and os.path.exists("temp_qrcode.png"):
            os.remove("temp_qrcode.png")
            
        return output_path

class Auth:
    @staticmethod
    def verify_admin(password):
        return password == "NaniGopal"
