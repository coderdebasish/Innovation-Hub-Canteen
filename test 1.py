from xml.dom.minidom import Document
import docx
import os

def rep_word(doc_name, old_word, new_word):
    # Open the Word document
    doc = docx.Document(doc_name)

    # Find and replace the old word with the new word
    for p in doc.paragraphs:
        if old_word in p.text:
            inline = p.runs
            # Loop through the runs in the paragraph
            for i in range(len(inline)):
                if old_word in inline[i].text:
                    text = inline[i].text.replace(old_word, new_word)
                    inline[i].text = text
    new_doc_name = "abcd.docx"
    doc.save(new_doc_name)

file = "Templetes//1.docx"
invno = "002"
rep_word(f"{file}", "111", f"INV{invno}")




















# Test the function with a sample Word document
# replace_word1("1QR.docx", "111", "001")
# replace_word1("abcd.docx", "112", "17.04.2023")
# # replace_word1("abcd.docx", "113", "Online")
# replace_word1("abcd.docx", "114", "Fish Meal")
# replace_word1("abcd.docx", "115", "2")
# replace_word1("abcd.docx", "116", "110")
# replace_word1("abcd.docx", "117", "2")
# replace_word1("abcd.docx", "118", "110")
# document = docx.Document("abcd.docx")
# document.add_picture("upi.jpg", "100", "100")
# document.save("abcd.docx")
# file = "abcd.docx"
# os.startfile(file,'print')
