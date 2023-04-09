import docx

def edit_text_box(input_file, output_file, text_box_index, new_text):
    # Load the input file
    doc = docx.Document(input_file)
    
    # Get the text box by index
    text_box = doc.element.body.getchildren()[text_box_index]
    
    # Replace the existing text with new text without changing the formatting
    p = text_box.paragraphs[0]
    p.clear()
    run = p.add_run(new_text)
    font = text_box.style.font
    run.font.bold = font.bold
    run.font.italic = font.italic
    run.font.underline = font.underline
    run.font.strike = font.strike
    run.font.color.rgb = font.color.rgb
    run.font.name = font.name
    run.font.size = font.size
    
    # Save the output file with a new name
    doc.save(output_file)


edit_text_box('input.docx', 'output.docx', 0, 'This is the new text.')
