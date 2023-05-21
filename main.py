from fileinput import filename
from glob import escape
import os
from timeit import repeat
from xmlrpc.client import TRANSPORT_ERROR
import datetime
from xml.dom.minidom import Document
import docx
import os
import qrcode
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

directory_path = 'Invoices'

def rep_word(doc_name, old_word, new_word):
    doc = docx.Document(doc_name)

    for p in doc.paragraphs:
        if old_word in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_word in inline[i].text:
                    text = inline[i].text.replace(old_word, new_word)
                    inline[i].text = text
    new_doc_name = f"{invno}.docx"
    doc.save(f"Invoices//{new_doc_name}")
def clear_screen():
    os.system('cls' if os.name == 'nt' else 'clear')
clear_screen()
escape_num = 0
repet = 0
discard = 0
esc = 0
items = 0
acer = 0
qwe = 0
generateinvoice2 = 0
generateinvoice3 = 0
generateinvoice4 = 0
generateinvoice5 = 0
generateinvoice6 = 0
generateinvoice7 = 0
generateinvoice8 = 0
while True:
    while True:
        # clear_screen()
        print("---------------------------------------------------------")
        print("             Welcome to Invoicing Software")
        print("---------------------------------------------------------")
        print("Choose numbers below \n1. Generate Invoice\n2. Reprint last Invoice\n3. Reprint Invoice\n4. View Bills\n5. Products Available\n6. Search Product\n7. Admin Panel")
        choice = input("Enter Your Choice: ")
        if (choice == "1" or choice == "2" or choice == "3" or choice == "4" or choice == "5" or choice == "6" or choice == "7"):
            break
        else:
            print("Invalid Input! Try Again")
            # Invoice Generation
    if choice == "1":
        while True: #item 1
            while True:
                print("which Product you want to ADD?")
                item_code1 = input("Enter Item Code(Example = 120) : ")
                if item_code1.isnumeric():
                    if len(item_code1) == 3:
                        break
                    else:
                        print("Invalid Item Code. Item code Should be 3 Digits")
                else:
                    print("Only numbers are allowed for Item code. Try again!")
            while True:
                if os.path.exists(f"Items\\{item_code1}.txt"):
                    break
                else:
                    print("No items in your Inventory with this product code.")
                    while True:
                        b = input("Do you want to add another product code ? (1 = Y& 0 = No): ")
                        if b == "1" or b == "0":
                            break
                        else:
                            print("Invalid Input! Try Again")
                            pass
                    if b == "1":
                        repet = 1
                        break
                    elif b == "0":
                        escape_num = 1
                        break
                    else:
                        print("Invalid Input! Try Again")
                    if escape_num == 1:
                        break
            if escape_num == 1:
                escape_num = 0
                break
            while True:
                if repet == 1:
                    repet = 0
                    break
                f = open(f"Items\\{item_code1}.txt", "r")
                itemname1 = f.readline()
                itemprice1 = f.readline()
                itemstock1 = f.readline()
                itemname1 = itemname1.strip()
                itemprice1 = itemprice1.strip()
                itemstock1 =itemstock1.strip()
                itemprice1 = int(itemprice1)
                itemstock1 = int(itemstock1)
                f.close()
                print(f"_ _ _ Product Details _ _ _\nItem Name = {itemname1}\nItem MRP/Unit = {itemprice1}\nItem Stock = {itemstock1}")
                while True:
                    item1qty = input("Enter Quantity : ")
                    if item1qty.isnumeric():
                        if int(item1qty) >= 0 and int(item1qty) <100:
                            item1qty = int(item1qty)
                            if item1qty <= itemstock1:
                                new_item_stock = (itemstock1 - item1qty)
                                f = open(f"Items\\{item_code1}.txt", "w")
                                f.write(f"{itemname1}\n{itemprice1}\n{new_item_stock}")
                                f.close()
                                print("New Item Stock = ", new_item_stock)
                                break
                            else:
                                print("You don't have enough stock to fulfill the order")
                                while True:
                                    c = input("Do you want to continue with this stock ? (1 = Y& 0 = No): ")
                                    if c == "1" or c == "0":
                                        break
                                    else:
                                        print("Invalid Input! Try Again")
                                        pass
                                if c == "1":
                                    pass
                                elif c == "0":
                                    acer = 1
                                    break
                                else:
                                    print("Invalid Input! Try Again")
                        else:
                            print("Invalid Item Code. Item Quantity Should be less than 2 Digits")
                    else:
                        print("Only numbers are allowed for Item code. Try again!")
                if acer == 1:
                    acer = 0
                    break
                print("Item Added")
                items = 1
                item1amount = item1qty * itemprice1
                print(f"Current Bill Amount is = {item1amount}")
                while True:
                    in1 = input("1. Add another item\n2. Generate Incoice\nEnter Your Choice : ")
                    if in1 == "1" or in1 == "2":
                        if in1 == "1":
                            add_item1 = 1
                            generateinvoice1 = 0
                            break
                        elif in1 == "2":
                            generateinvoice1 = 1
                            add_item1 = 0
                            break
                        else:
                            print("Invalid Input! Please Try Again.")
                    else:
                        print("Invalid Input! Please Try Again.")
                if generateinvoice1 == 1:
                    while True:
                        paymode = input("Payment Mode:\n1. Cash Payment\n2. Upi Payment\n Enter your Choice : ")
                        if paymode == "1" or paymode == "2":
                            if paymode == "1":
                                file = "Templetes//1.docx"
                                break
                            elif paymode == "2":
                                file = "Templetes//1QR.docx"
                                link = f"upi://pay?pa=rajlakshmi.mohanty@ybl&pn=abcd&am={item1amount}&tn=Tasty Confectionary&cu=INR"
                                img = qrcode.make(link)
                                img.save("qrcode.png")
                                break
                        else:
                            print("Invalid Input! Please Try Again.")
                    files = os.listdir(directory_path)
                    num_files = len(files)
                    invno = num_files + 1
                    print("Invoice Number is = ", invno)
                    now = datetime.datetime.now()
                    current_datetime = now.strftime("%d-%m-%Y        Time - %H:%M:%S")
                    rep_word(f"{file}", "111", f"INV {invno}")
                    rep_word(f"Invoices//{invno}.docx", "112", f"{current_datetime}")
                    rep_word(f"Invoices//{invno}.docx", "114", f"{itemname1}")
                    rep_word(f"Invoices//{invno}.docx", "115", f"{item1qty}")
                    rep_word(f"Invoices//{invno}.docx", "116", f"{item1amount}")
                    rep_word(f"Invoices//{invno}.docx", "117", f"{item1qty}")
                    rep_word(f"Invoices//{invno}.docx", "118", f"{item1amount}")
                    if paymode == "2":
                        document = Document(f"Invoices//{invno}.docx")
                        paragraph = document.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        run.add_picture('qrcode.png', width=Inches(0.85), height=Inches(0.85))
                        document.save(f"Invoices//{invno}.docx")
                    while True:
                        if os.path.exists(f"Invoices//{invno}.docx"):
                            break
                    file = f"Invoices\\{invno}.docx"
                    os.startfile(file)
                    esc = 1
                    break
                elif add_item1 == 1: # item 2
                    add_item1 = 0
                    while True:
                        while True:
                            print("which Product you want to ADD?")
                            item_code2 = input("Enter Item Code(Example = 120) : ")
                            if item_code2.isnumeric():
                                if len(item_code2) == 3:
                                    break
                                else:
                                    print("Invalid Item Code. Item code Should be 3 Digits")
                            else:
                                print("Only numbers are allowed for Item code. Try again!")
                        while True:
                            if os.path.exists(f"Items\\{item_code2}.txt"):
                                break
                            else:
                                print("No items in your Inventory with this product code.")
                                while True:
                                    repet = 0
                                    escape_num = 0
                                    b = input("Do you want to add another product code ? (1 = Y& 0 = No): ")
                                    if b == "1" or b == "0":
                                        break
                                    else:
                                        print("Invalid Input! Try Again")
                                        pass
                                if b == "1":
                                    repet = 1
                                    break
                                elif b == "0":
                                    escape_num = 1
                                    break
                                else:
                                    print("Invalid Input! Try Again")
                                if escape_num == 1:
                                    break
                        if escape_num == 1:
                            escape_num = 0
                            break
                        while True:
                            if repet == 1:
                                repet = 0
                                break
                            f = open(f"Items\\{item_code2}.txt", "r")
                            itemname2 = f.readline()
                            itemprice2 = f.readline()
                            itemstock2 = f.readline()
                            itemname2 = itemname2.strip()
                            itemprice2 = itemprice2.strip()
                            itemstock2 =itemstock2.strip()
                            itemprice2 = int(itemprice2)
                            itemstock2 = int(itemstock2)
                            f.close()
                            print(f"_ _ _ Product Details _ _ _\nItem Name = {itemname2}\nItem MRP/Unit = {itemprice2}\nItem Stock = {itemstock2}")
                            while True:
                                item2qty = input("Enter Quantity : ")
                                if item2qty.isnumeric():
                                    if int(item2qty) >= 0 and int(item2qty) <100:
                                        item2qty = int(item2qty)
                                        if item2qty <= itemstock2:
                                            new_item_stock = (itemstock2 - item2qty)
                                            f = open(f"Items\\{item_code2}.txt", "w")
                                            f.write(f"{itemname2}\n{itemprice2}\n{new_item_stock}")
                                            f.close()
                                            print("New Item Stock = ", new_item_stock)
                                            break
                                        else:
                                            print("You don't have enough stock to fulfill the order")
                                            while True:
                                                c = input("Do you want to continue with this stock ? (1 = Y& 0 = No): ")
                                                if c == "1" or c == "0":
                                                    break
                                                else:
                                                    print("Invalid Input! Try Again")
                                                    pass
                                            if c == "1":
                                                pass
                                            elif c == "0":
                                                acer = 1
                                                break
                                            else:
                                                print("Invalid Input! Try Again")
                                    else:
                                        print("Invalid Item Code. Item Quantity Should be less than 2 Digits")
                                else:
                                    print("Only numbers are allowed for Item code. Try again!")
                            if acer == 1:
                                acer = 0
                                break

                            print("Item 2 Added")
                            items = 2
                            item1amount = item1qty * itemprice1
                            item2amount = item2qty * itemprice2
                            item2total = item1amount+item2amount
                            item2qtytotal = item1qty + item2qty

                            print(f"Current Bill Amount is = {item2total}")
                            while True:
                                in2 = input("1. Add another item\n2. Generate Incoice\n Enter your Choice : ")
                                if in2 == "1" or in2 == "2":
                                    if in2 == "1":
                                        add_item2 = 1
                                        break
                                    elif in2 == "2":
                                        generateinvoice2 = 1
                                        break
                                    else:
                                        print("Invalid Input! Please Try Again.")
                                else:
                                    print("Invalid Input! Please Try Again.")

                            if generateinvoice2 == 1:
                                generateinvoice2 = 0
                                while True:
                                    paymode2 = input("Payment Mode:\n1. Cash Payment\n2. Upi Payment\n Enter your Choice : ")
                                    if paymode2 == "1" or paymode2 == "2":
                                        if paymode2 == "1":
                                            file2 = "Templetes//2.docx"
                                            break
                                        elif paymode2 == "2":
                                            file2 = "Templetes//2QR.docx"
                                            link = f"upi://pay?pa=rajlakshmi.mohanty@ybl&pn=abcd&am={item2total}&tn=Tasty Confectionary&cu=INR"
                                            img = qrcode.make(link)
                                            img.save("qrcode.png")
                                            break
                                    else:
                                        print("Invalid Input! Please Try Again.")
                                files = os.listdir(directory_path)
                                num_files = len(files)
                                invno = num_files + 1
                                print("Invoice Number is = ", invno)
                                now = datetime.datetime.now()
                                current_datetime = now.strftime("%d-%m-%Y        Time - %H:%M:%S")
                                rep_word(f"{file2}", "111", f"INV {invno}")
                                rep_word(f"Invoices//{invno}.docx", "112", f"{current_datetime}")
                                rep_word(f"Invoices//{invno}.docx", "114", f"{itemname1}")
                                rep_word(f"Invoices//{invno}.docx", "115", f"{item1qty}")
                                rep_word(f"Invoices//{invno}.docx", "116", f"{item1amount}")
                                rep_word(f"Invoices//{invno}.docx", "117", f"{itemname2}")
                                rep_word(f"Invoices//{invno}.docx", "118", f"{item2qty}")
                                rep_word(f"Invoices//{invno}.docx", "119", f"{item2amount}")
                                rep_word(f"Invoices//{invno}.docx", "55", f"{item2qtytotal}")
                                rep_word(f"Invoices//{invno}.docx", "99", f"{item2total}")
                                # print(item2total, "=" , item2qtytotal)
                                if paymode2 == "2":
                                    document = Document(f"Invoices//{invno}.docx")
                                    paragraph = document.add_paragraph()
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = paragraph.add_run()
                                    run.add_picture('qrcode.png', width=Inches(0.8), height=Inches(0.8))

                                    document.save(f"Invoices//{invno}.docx")
                                while True:
                                    if os.path.exists(f"Invoices//{invno}.docx"):
                                        break
                                file = f"Invoices\\{invno}.docx"
                                os.startfile(file)
                                esc = 1
                                break
                            elif add_item2 == 1: # item 3
                                add_item2 = 0
                                while True:
                                    while True:
                                        print("which Product you want to ADD?")
                                        item_code3 = input("Enter Item Code(Example = 120) : ")
                                        if item_code3.isnumeric():
                                            if len(item_code3) == 3:
                                                break
                                            else:
                                                print("Invalid Item Code. Item code Should be 3 Digits")
                                        else:
                                            print("Only numbers are allowed for Item code. Try again!")
                                    while True:
                                        if os.path.exists(f"Items\\{item_code3}.txt"):
                                            break
                                        else:
                                            print("No items in your Inventory with this product code.")
                                            while True:
                                                repet = 0
                                                escape_num = 0
                                                b = input("Do you want to add another product code ? (1 = Y& 0 = No): ")
                                                if b == "1" or b == "0":
                                                    break
                                                else:
                                                    print("Invalid Input! Try Again")
                                                    pass
                                            if b == "1":
                                                repet = 1
                                                break
                                            elif b == "0":
                                                escape_num = 1
                                                break
                                            else:
                                                print("Invalid Input! Try Again")
                                            if escape_num == 1:
                                                break
                                    if escape_num == 1:
                                        escape_num = 0
                                        break
                                    while True:
                                        if repet == 1:
                                            repet = 0
                                            break
                                        f = open(f"Items\\{item_code3}.txt", "r")
                                        itemname3 = f.readline()
                                        itemprice3 = f.readline()
                                        itemstock3 = f.readline()
                                        itemname3 = itemname3.strip()
                                        itemprice3 = itemprice3.strip()
                                        itemstock3 =itemstock3.strip()
                                        itemprice3 = int(itemprice3)
                                        itemstock3 = int(itemstock3)
                                        f.close()
                                        print(f"_ _ _ Product Details _ _ _\nItem Name = {itemname3}\nItem MRP/Unit = {itemprice3}\nItem Stock = {itemstock3}")
                                        while True:
                                            item3qty = input("Enter Quantity : ")
                                            if item3qty.isnumeric():
                                                if int(item3qty) >= 0 and int(item3qty) <100:
                                                    item3qty = int(item3qty)
                                                    if item3qty <= itemstock3:
                                                        new_item_stock = (itemstock3 - item3qty)
                                                        f = open(f"Items\\{item_code3}.txt", "w")
                                                        f.write(f"{itemname3}\n{itemprice3}\n{new_item_stock}")
                                                        f.close()
                                                        print("New Item Stock = ", new_item_stock)
                                                        break
                                                    else:
                                                        print("You don't have enough stock to fulfill the order")
                                                        while True:
                                                            c = input("Do you want to continue with this stock ? (1 = Y& 0 = No): ")
                                                            if c == "1" or c == "0":
                                                                break
                                                            else:
                                                                print("Invalid Input! Try Again")
                                                                pass
                                                        if c == "1":
                                                            pass
                                                        elif c == "0":
                                                            acer = 1
                                                            break
                                                        else:
                                                            print("Invalid Input! Try Again")
                                                else:
                                                    print("Invalid Item Code. Item Quantity Should be less than 2 Digits")
                                            else:
                                                print("Only numbers are allowed for Item code. Try again!")
                                        if acer == 1:
                                            acer = 0
                                            break

                                        print("Item 3 Added")
                                        items = 3
                                        item1amount = item1qty * itemprice1
                                        item2amount = item2qty * itemprice2
                                        item3amount = item3qty * itemprice3
                                        item3total = item1amount+item2amount+item3amount
                                        item3qtytotal = item1qty + item2qty + item3qty

                                        print(f"Current Bill Amount is = {item3total}")
                                        while True:
                                            in3 = input("1. Add another item\n2. Generate Incoice\n Enter your Choice : ")
                                            if in3 == "1" or in3 == "2":
                                                if in3 == "1":
                                                    add_item4 = 1
                                                    break
                                                elif in3 == "2":
                                                    generateinvoice3 = 1
                                                    break
                                                else:
                                                    print("Invalid Input! Please Try Again.")
                                            else:
                                                print("Invalid Input! Please Try Again.")

                                        if generateinvoice3 == 1:
                                            generateinvoice3 = 0
                                            while True:
                                                paymode2 = input("Payment Mode:\n1. Cash Payment\n2. Upi Payment\n Enter your Choice : ")
                                                if paymode2 == "1" or paymode2 == "2":
                                                    if paymode2 == "1":
                                                        file2 = "Templetes//3.docx"
                                                        break
                                                    elif paymode2 == "2":
                                                        file2 = "Templetes//3QR.docx"
                                                        link = f"upi://pay?pa=rajlakshmi.mohanty@ybl&pn=abcd&am={item3total}&tn=Tasty Confectionary&cu=INR"
                                                        img = qrcode.make(link)
                                                        img.save("qrcode.png")
                                                        break
                                                else:
                                                    print("Invalid Input! Please Try Again.")
                                            files = os.listdir(directory_path)
                                            num_files = len(files)
                                            invno = num_files + 1
                                            print("Invoice Number is = ", invno)
                                            now = datetime.datetime.now()
                                            current_datetime = now.strftime("%d-%m-%Y        Time - %H:%M:%S")
                                            rep_word(f"{file2}", "111", f"INV {invno}")
                                            rep_word(f"Invoices//{invno}.docx", "112", f"{current_datetime}")
                                            rep_word(f"Invoices//{invno}.docx", "114", f"{itemname1}")
                                            rep_word(f"Invoices//{invno}.docx", "115", f"{item1qty}")
                                            rep_word(f"Invoices//{invno}.docx", "116", f"{item1amount}")
                                            rep_word(f"Invoices//{invno}.docx", "117", f"{itemname2}")
                                            rep_word(f"Invoices//{invno}.docx", "118", f"{item2qty}")
                                            rep_word(f"Invoices//{invno}.docx", "119", f"{item2amount}")
                                            rep_word(f"Invoices//{invno}.docx", "120", f"{itemname3}")
                                            rep_word(f"Invoices//{invno}.docx", "121", f"{item3qty}")
                                            rep_word(f"Invoices//{invno}.docx", "122", f"{item3amount}")
                                            rep_word(f"Invoices//{invno}.docx", "55", f"{item3qtytotal}")
                                            rep_word(f"Invoices//{invno}.docx", "99", f"{item3total}")
                                            # print(item2total, "=" , item2qtytotal)
                                            if paymode2 == "2":
                                                document = Document(f"Invoices//{invno}.docx")
                                                paragraph = document.add_paragraph()
                                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                run = paragraph.add_run()
                                                run.add_picture('qrcode.png', width=Inches(0.8), height=Inches(0.8))

                                                document.save(f"Invoices//{invno}.docx")
                                            while True:
                                                if os.path.exists(f"Invoices//{invno}.docx"):
                                                    break
                                            file = f"Invoices\\{invno}.docx"
                                            os.startfile(file)
                                            esc = 1
                                            break
                                        elif add_item4 == 1: # item 4
                                            add_item4 = 0
                                            while True:
                                                while True:
                                                    print("which Product you want to ADD?")
                                                    item_code4 = input("Enter Item Code(Example = 120) : ")
                                                    if item_code4.isnumeric():
                                                        if len(item_code4) == 3:
                                                            break
                                                        else:
                                                            print("Invalid Item Code. Item code Should be 3 Digits")
                                                    else:
                                                        print("Only numbers are allowed for Item code. Try again!")
                                                while True:
                                                    if os.path.exists(f"Items\\{item_code4}.txt"):
                                                        break
                                                    else:
                                                        print("No items in your Inventory with this product code.")
                                                        while True:
                                                            repet = 0
                                                            escape_num = 0
                                                            b = input("Do you want to add another product code ? (1 = Y& 0 = No): ")
                                                            if b == "1" or b == "0":
                                                                break
                                                            else:
                                                                print("Invalid Input! Try Again")
                                                                pass
                                                        if b == "1":
                                                            repet = 1
                                                            break
                                                        elif b == "0":
                                                            escape_num = 1
                                                            break
                                                        else:
                                                            print("Invalid Input! Try Again")
                                                        if escape_num == 1:
                                                            break
                                                if escape_num == 1:
                                                    escape_num = 0
                                                    break
                                                while True:
                                                    if repet == 1:
                                                        repet = 0
                                                        break
                                                    f = open(f"Items\\{item_code4}.txt", "r")
                                                    itemname4 = f.readline()
                                                    itemprice4 = f.readline()
                                                    itemstock4 = f.readline()
                                                    itemname4 = itemname4.strip()
                                                    itemprice4 = itemprice4.strip()
                                                    itemstock4 =itemstock4.strip()
                                                    itemprice4 = int(itemprice4)
                                                    itemstock4 = int(itemstock4)
                                                    f.close()
                                                    print(f"_ _ _ Product Details _ _ _\nItem Name = {itemname4}\nItem MRP/Unit = {itemprice4}\nItem Stock = {itemstock4}")
                                                    while True:
                                                        item4qty = input("Enter Quantity : ")
                                                        if item4qty.isnumeric():
                                                            if int(item4qty) >= 0 and int(item4qty) <100:
                                                                item4qty = int(item4qty)
                                                                if item4qty <= itemstock4:
                                                                    new_item_stock = (itemstock4 - item4qty)
                                                                    f = open(f"Items\\{item_code4}.txt", "w")
                                                                    f.write(f"{itemname4}\n{itemprice4}\n{new_item_stock}")
                                                                    f.close()
                                                                    print("New Item Stock = ", new_item_stock)
                                                                    break
                                                                else:
                                                                    print("You don't have enough stock to fulfill the order")
                                                                    while True:
                                                                        c = input("Do you want to continue with this stock ? (1 = Y& 0 = No): ")
                                                                        if c == "1" or c == "0":
                                                                            break
                                                                        else:
                                                                            print("Invalid Input! Try Again")
                                                                            pass
                                                                    if c == "1":
                                                                        pass
                                                                    elif c == "0":
                                                                        acer = 1
                                                                        break
                                                                    else:
                                                                        print("Invalid Input! Try Again")
                                                            else:
                                                                print("Invalid Item Code. Item Quantity Should be less than 2 Digits")
                                                        else:
                                                            print("Only numbers are allowed for Item code. Try again!")
                                                    if acer == 1:
                                                        acer = 0
                                                        break

                                                    print("Item 4 Added")
                                                    items = 4
                                                    item1amount = item1qty * itemprice1
                                                    item2amount = item2qty * itemprice2
                                                    item3amount = item3qty * itemprice3
                                                    item4amount = item4qty * itemprice4
                                                    item4total = item1amount+item2amount+item3amount+item4amount
                                                    item4qtytotal = item1qty + item2qty + item3qty + item4qty

                                                    print(f"Current Bill Amount is = {item4total}")
                                                    while True:
                                                        in3 = input("1. Add another item\n2. Generate Incoice\n Enter your Choice : ")
                                                        if in3 == "1" or in3 == "2":
                                                            if in3 == "1":
                                                                add_item5 = 1
                                                                break
                                                            elif in3 == "2":
                                                                generateinvoice4 = 1
                                                                break
                                                            else:
                                                                print("Invalid Input! Please Try Again.")
                                                        else:
                                                            print("Invalid Input! Please Try Again.")

                                                    if generateinvoice4 == 1:
                                                        generateinvoice4 = 0
                                                        while True:
                                                            paymode3 = input("Payment Mode:\n1. Cash Payment\n2. Upi Payment\n Enter your Choice : ")
                                                            if paymode3 == "1" or paymode3 == "2":
                                                                if paymode3 == "1":
                                                                    file2 = "Templetes//4.docx"
                                                                    break
                                                                elif paymode3 == "2":
                                                                    file2 = "Templetes//4QR.docx"
                                                                    link = f"upi://pay?pa=rajlakshmi.mohanty@ybl&pn=abcd&am={item4total}&tn=Tasty Confectionary&cu=INR"
                                                                    img = qrcode.make(link)
                                                                    img.save("qrcode.png")
                                                                    break
                                                            else:
                                                                print("Invalid Input! Please Try Again.")
                                                        files = os.listdir(directory_path)
                                                        num_files = len(files)
                                                        invno = num_files + 1
                                                        print("Invoice Number is = ", invno)
                                                        now = datetime.datetime.now()
                                                        current_datetime = now.strftime("%d-%m-%Y        Time - %H:%M:%S")
                                                        rep_word(f"{file2}", "111", f"INV {invno}")
                                                        rep_word(f"Invoices//{invno}.docx", "112", f"{current_datetime}")
                                                        rep_word(f"Invoices//{invno}.docx", "114", f"{itemname1}")
                                                        rep_word(f"Invoices//{invno}.docx", "115", f"{item1qty}")
                                                        rep_word(f"Invoices//{invno}.docx", "116", f"{item1amount}")
                                                        rep_word(f"Invoices//{invno}.docx", "117", f"{itemname2}")
                                                        rep_word(f"Invoices//{invno}.docx", "118", f"{item2qty}")
                                                        rep_word(f"Invoices//{invno}.docx", "66", f"{item2amount}")
                                                        rep_word(f"Invoices//{invno}.docx", "120", f"{itemname3}")
                                                        rep_word(f"Invoices//{invno}.docx", "121", f"{item3qty}")
                                                        rep_word(f"Invoices//{invno}.docx", "122", f"{item3amount}")
                                                        rep_word(f"Invoices//{invno}.docx", "123", f"{itemname4}")
                                                        rep_word(f"Invoices//{invno}.docx", "124", f"{item4qty}")
                                                        rep_word(f"Invoices//{invno}.docx", "125", f"{item4amount}")
                                                        rep_word(f"Invoices//{invno}.docx", "55", f"{item4qtytotal}")
                                                        rep_word(f"Invoices//{invno}.docx", "99", f"{item4total}")
                                                        # print(item2total, "=" , item2qtytotal)
                                                        if paymode3 == "2":
                                                            document = Document(f"Invoices//{invno}.docx")
                                                            paragraph = document.add_paragraph()
                                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                            run = paragraph.add_run()
                                                            run.add_picture('qrcode.png', width=Inches(0.8), height=Inches(0.8))

                                                            document.save(f"Invoices//{invno}.docx")
                                                        while True:
                                                            if os.path.exists(f"Invoices//{invno}.docx"):
                                                                break
                                                        file = f"Invoices\\{invno}.docx"
                                                        os.startfile(file)
                                                        esc = 1
                                                        break
                                                    elif add_item5 == 1: # item 5
                                                        add_item5 = 0
                                                        while True:
                                                            while True:
                                                                print("which Product you want to ADD?")
                                                                item_code5 = input("Enter Item Code(Example = 120) : ")
                                                                if item_code5.isnumeric():
                                                                    if len(item_code5) == 3:
                                                                        break
                                                                    else:
                                                                        print("Invalid Item Code. Item code Should be 3 Digits")
                                                                else:
                                                                    print("Only numbers are allowed for Item code. Try again!")
                                                            while True:
                                                                if os.path.exists(f"Items\\{item_code5}.txt"):
                                                                    break
                                                                else:
                                                                    print("No items in your Inventory with this product code.")
                                                                    while True:
                                                                        repet = 0
                                                                        escape_num = 0
                                                                        b = input("Do you want to add another product code ? (1 = Y& 0 = No): ")
                                                                        if b == "1" or b == "0":
                                                                            break
                                                                        else:
                                                                            print("Invalid Input! Try Again")
                                                                            pass
                                                                    if b == "1":
                                                                        repet = 1
                                                                        break
                                                                    elif b == "0":
                                                                        escape_num = 1
                                                                        break
                                                                    else:
                                                                        print("Invalid Input! Try Again")
                                                                    if escape_num == 1:
                                                                        break
                                                            if escape_num == 1:
                                                                escape_num = 0
                                                                break
                                                            while True:
                                                                if repet == 1:
                                                                    repet = 0
                                                                    break
                                                                f = open(f"Items\\{item_code5}.txt", "r")
                                                                itemname5 = f.readline()
                                                                itemprice5 = f.readline()
                                                                itemstock5 = f.readline()
                                                                itemname5 = itemname5.strip()
                                                                itemprice5 = itemprice5.strip()
                                                                itemstock5 =itemstock5.strip()
                                                                itemprice5 = int(itemprice5)
                                                                itemstock5 = int(itemstock5)
                                                                f.close()
                                                                print(f"_ _ _ Product Details _ _ _\nItem Name = {itemname5}\nItem MRP/Unit = {itemprice5}\nItem Stock = {itemstock5}")
                                                                while True:
                                                                    item5qty = input("Enter Quantity : ")
                                                                    if item5qty.isnumeric():
                                                                        if int(item5qty) >= 0 and int(item5qty) <100:
                                                                            item5qty = int(item5qty)
                                                                            if item5qty <= itemstock5:
                                                                                new_item_stock = (itemstock5 - item5qty)
                                                                                f = open(f"Items\\{item_code5}.txt", "w")
                                                                                f.write(f"{itemname5}\n{itemprice5}\n{new_item_stock}")
                                                                                f.close()
                                                                                print("New Item Stock = ", new_item_stock)
                                                                                break
                                                                            else:
                                                                                print("You don't have enough stock to fulfill the order")
                                                                                while True:
                                                                                    c = input("Do you want to continue with this stock ? (1 = Y& 0 = No): ")
                                                                                    if c == "1" or c == "0":
                                                                                        break
                                                                                    else:
                                                                                        print("Invalid Input! Try Again")
                                                                                        pass
                                                                                if c == "1":
                                                                                    pass
                                                                                elif c == "0":
                                                                                    acer = 1
                                                                                    break
                                                                                else:
                                                                                    print("Invalid Input! Try Again")
                                                                        else:
                                                                            print("Invalid Item Code. Item Quantity Should be less than 2 Digits")
                                                                    else:
                                                                        print("Only numbers are allowed for Item code. Try again!")
                                                                if acer == 1:
                                                                    acer = 0
                                                                    break

                                                                print("Item 5 Added")
                                                                items = 5
                                                                item1amount = item1qty * itemprice1
                                                                item2amount = item2qty * itemprice2
                                                                item3amount = item3qty * itemprice3
                                                                item4amount = item4qty * itemprice4
                                                                item5amount = item5qty * itemprice5
                                                                item5total = item1amount+item2amount+item3amount+item4amount+item5amount
                                                                item5qtytotal = item1qty + item2qty + item3qty + item4qty + item5qty

                                                                print(f"Current Bill Amount is = {item5total}")
                                                                while True:
                                                                    in3 = input("1. Add another item\n2. Generate Incoice\n Enter your Choice : ")
                                                                    if in3 == "1" or in3 == "2":
                                                                        if in3 == "1":
                                                                            add_item6 = 1
                                                                            break
                                                                        elif in3 == "2":
                                                                            generateinvoice5 = 1
                                                                            break
                                                                        else:
                                                                            print("Invalid Input! Please Try Again.")
                                                                    else:
                                                                        print("Invalid Input! Please Try Again.")

                                                                if generateinvoice5 == 1:
                                                                    generateinvoice5 = 0
                                                                    while True:
                                                                        paymode3 = input("Payment Mode:\n1. Cash Payment\n2. Upi Payment\n Enter your Choice : ")
                                                                        if paymode3 == "1" or paymode3 == "2":
                                                                            if paymode3 == "1":
                                                                                file2 = "Templetes//5.docx"
                                                                                break
                                                                            elif paymode3 == "2":
                                                                                file2 = "Templetes//5QR.docx"
                                                                                link = f"upi://pay?pa=rajlakshmi.mohanty@ybl&pn=abcd&am={item5total}&tn=Tasty Confectionary&cu=INR"
                                                                                img = qrcode.make(link)
                                                                                img.save("qrcode.png")
                                                                                break
                                                                        else:
                                                                            print("Invalid Input! Please Try Again.")
                                                                    files = os.listdir(directory_path)
                                                                    num_files = len(files)
                                                                    invno = num_files + 1
                                                                    print("Invoice Number is = ", invno)
                                                                    now = datetime.datetime.now()
                                                                    current_datetime = now.strftime("%d-%m-%Y        Time - %H:%M:%S")
                                                                    rep_word(f"{file2}", "111", f"INV {invno}")
                                                                    rep_word(f"Invoices//{invno}.docx", "112", f"{current_datetime}")
                                                                    rep_word(f"Invoices//{invno}.docx", "114", f"{itemname1}")
                                                                    rep_word(f"Invoices//{invno}.docx", "115", f"{item1qty}")
                                                                    rep_word(f"Invoices//{invno}.docx", "116", f"{item1amount}")
                                                                    rep_word(f"Invoices//{invno}.docx", "117", f"{itemname2}")
                                                                    rep_word(f"Invoices//{invno}.docx", "118", f"{item2qty}")
                                                                    rep_word(f"Invoices//{invno}.docx", "119", f"{item2amount}")
                                                                    rep_word(f"Invoices//{invno}.docx", "120", f"{itemname3}")
                                                                    rep_word(f"Invoices//{invno}.docx", "121", f"{item3qty}")
                                                                    rep_word(f"Invoices//{invno}.docx", "122", f"{item3amount}")
                                                                    rep_word(f"Invoices//{invno}.docx", "123", f"{itemname4}")
                                                                    rep_word(f"Invoices//{invno}.docx", "124", f"{item4qty}")
                                                                    rep_word(f"Invoices//{invno}.docx", "125", f"{item4amount}")
                                                                    rep_word(f"Invoices//{invno}.docx", "126", f"{itemname5}")
                                                                    rep_word(f"Invoices//{invno}.docx", "127", f"{item5qty}")
                                                                    rep_word(f"Invoices//{invno}.docx", "128", f"{item5amount}")
                                                                    rep_word(f"Invoices//{invno}.docx", "55", f"{item5qtytotal}")
                                                                    rep_word(f"Invoices//{invno}.docx", "99", f"{item5total}")
                                                                    # print(item2total, "=" , item2qtytotal)
                                                                    if paymode3 == "2":
                                                                        document = Document(f"Invoices//{invno}.docx")
                                                                        paragraph = document.add_paragraph()
                                                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                                        run = paragraph.add_run()
                                                                        run.add_picture('qrcode.png', width=Inches(0.8), height=Inches(0.8))

                                                                        document.save(f"Invoices//{invno}.docx")
                                                                    while True:
                                                                        if os.path.exists(f"Invoices//{invno}.docx"):
                                                                            break
                                                                    file = f"Invoices\\{invno}.docx"
                                                                    os.startfile(file)
                                                                    esc = 1
                                                                    break
                                                                elif add_item6 == 1: # item 6
                                                                    add_item6 = 0
                                                                    while True:
                                                                        while True:
                                                                            print("which Product you want to ADD?")
                                                                            item_code6 = input("Enter Item Code(Example = 120) : ")
                                                                            if item_code6.isnumeric():
                                                                                if len(item_code6) == 3:
                                                                                    break
                                                                                else:
                                                                                    print("Invalid Item Code. Item code Should be 3 Digits")
                                                                            else:
                                                                                print("Only numbers are allowed for Item code. Try again!")
                                                                        while True:
                                                                            if os.path.exists(f"Items\\{item_code6}.txt"):
                                                                                break
                                                                            else:
                                                                                print("No items in your Inventory with this product code.")
                                                                                while True:
                                                                                    repet = 0
                                                                                    escape_num = 0
                                                                                    b = input("Do you want to add another product code ? (1 = Y& 0 = No): ")
                                                                                    if b == "1" or b == "0":
                                                                                        break
                                                                                    else:
                                                                                        print("Invalid Input! Try Again")
                                                                                        pass
                                                                                if b == "1":
                                                                                    repet = 1
                                                                                    break
                                                                                elif b == "0":
                                                                                    escape_num = 1
                                                                                    break
                                                                                else:
                                                                                    print("Invalid Input! Try Again")
                                                                                if escape_num == 1:
                                                                                    break
                                                                        if escape_num == 1:
                                                                            escape_num = 0
                                                                            break
                                                                        while True:
                                                                            if repet == 1:
                                                                                repet = 0
                                                                                break
                                                                            f = open(f"Items\\{item_code6}.txt", "r")
                                                                            itemname6 = f.readline()
                                                                            itemprice6 = f.readline()
                                                                            itemstock6 = f.readline()
                                                                            itemname6 = itemname6.strip()
                                                                            itemprice6 = itemprice6.strip()
                                                                            itemstock6 =itemstock6.strip()
                                                                            itemprice6 = int(itemprice6)
                                                                            itemstock6 = int(itemstock6)
                                                                            f.close()
                                                                            print(f"_ _ _ Product Details _ _ _\nItem Name = {itemname6}\nItem MRP/Unit = {itemprice6}\nItem Stock = {itemstock6}")
                                                                            while True:
                                                                                item6qty = input("Enter Quantity : ")
                                                                                if item6qty.isnumeric():
                                                                                    if int(item6qty) >= 0 and int(item6qty) <100:
                                                                                        item6qty = int(item6qty)
                                                                                        if item6qty <= itemstock6:
                                                                                            new_item_stock = (itemstock6 - item6qty)
                                                                                            f = open(f"Items\\{item_code6}.txt", "w")
                                                                                            f.write(f"{itemname6}\n{itemprice6}\n{new_item_stock}")
                                                                                            f.close()
                                                                                            print("New Item Stock = ", new_item_stock)
                                                                                            break
                                                                                        else:
                                                                                            print("You don't have enough stock to fulfill the order")
                                                                                            while True:
                                                                                                c = input("Do you want to continue with this stock ? (1 = Y& 0 = No): ")
                                                                                                if c == "1" or c == "0":
                                                                                                    break
                                                                                                else:
                                                                                                    print("Invalid Input! Try Again")
                                                                                                    pass
                                                                                            if c == "1":
                                                                                                pass
                                                                                            elif c == "0":
                                                                                                acer = 1
                                                                                                break
                                                                                            else:
                                                                                                print("Invalid Input! Try Again")
                                                                                    else:
                                                                                        print("Invalid Item Code. Item Quantity Should be less than 2 Digits")
                                                                                else:
                                                                                    print("Only numbers are allowed for Item code. Try again!")
                                                                            if acer == 1:
                                                                                acer = 0
                                                                                break

                                                                            print("Item 6 Added")
                                                                            items = 6
                                                                            item1amount = item1qty * itemprice1
                                                                            item2amount = item2qty * itemprice2
                                                                            item3amount = item3qty * itemprice3
                                                                            item4amount = item4qty * itemprice4
                                                                            item5amount = item5qty * itemprice5
                                                                            item6amount = item6qty * itemprice6
                                                                            item6total = item1amount+item2amount+item3amount+item4amount+item5amount+item6amount
                                                                            item6qtytotal = item1qty + item2qty + item3qty + item4qty + item5qty + item6qty

                                                                            print(f"Current Bill Amount is = {item6total}")
                                                                            while True:
                                                                                in3 = input("1. Add another item\n2. Generate Incoice\n Enter your Choice : ")
                                                                                if in3 == "1" or in3 == "2":
                                                                                    if in3 == "1":
                                                                                        add_item7 = 1
                                                                                        break
                                                                                    elif in3 == "2":
                                                                                        generateinvoice6 = 1
                                                                                        break
                                                                                    else:
                                                                                        print("Invalid Input! Please Try Again.")
                                                                                else:
                                                                                    print("Invalid Input! Please Try Again.")

                                                                            if generateinvoice6 == 1:
                                                                                generateinvoice6 = 0
                                                                                while True:
                                                                                    paymode3 = input("Payment Mode:\n1. Cash Payment\n2. Upi Payment\n Enter your Choice : ")
                                                                                    if paymode3 == "1" or paymode3 == "2":
                                                                                        if paymode3 == "1":
                                                                                            file2 = "Templetes//6.docx"
                                                                                            break
                                                                                        elif paymode3 == "2":
                                                                                            file2 = "Templetes//6QR.docx"
                                                                                            link = f"upi://pay?pa=rajlakshmi.mohanty@ybl&pn=abcd&am={item6total}&tn=Tasty Confectionary&cu=INR"
                                                                                            img = qrcode.make(link)
                                                                                            img.save("qrcode.png")
                                                                                            break
                                                                                    else:
                                                                                        print("Invalid Input! Please Try Again.")
                                                                                files = os.listdir(directory_path)
                                                                                num_files = len(files)
                                                                                invno = num_files + 1
                                                                                print("Invoice Number is = ", invno)
                                                                                now = datetime.datetime.now()
                                                                                current_datetime = now.strftime("%d-%m-%Y        Time - %H:%M:%S")
                                                                                rep_word(f"{file2}", "111", f"INV {invno}")
                                                                                rep_word(f"Invoices//{invno}.docx", "112", f"{current_datetime}")
                                                                                rep_word(f"Invoices//{invno}.docx", "114", f"{itemname1}")
                                                                                rep_word(f"Invoices//{invno}.docx", "115", f"{item1qty}")
                                                                                rep_word(f"Invoices//{invno}.docx", "116", f"{item1amount}")
                                                                                rep_word(f"Invoices//{invno}.docx", "117", f"{itemname2}")
                                                                                rep_word(f"Invoices//{invno}.docx", "118", f"{item2qty}")
                                                                                rep_word(f"Invoices//{invno}.docx", "119", f"{item2amount}")
                                                                                rep_word(f"Invoices//{invno}.docx", "120", f"{itemname3}")
                                                                                rep_word(f"Invoices//{invno}.docx", "121", f"{item3qty}")
                                                                                rep_word(f"Invoices//{invno}.docx", "122", f"{item3amount}")
                                                                                rep_word(f"Invoices//{invno}.docx", "123", f"{itemname4}")
                                                                                rep_word(f"Invoices//{invno}.docx", "124", f"{item4qty}")
                                                                                rep_word(f"Invoices//{invno}.docx", "125", f"{item4amount}")
                                                                                rep_word(f"Invoices//{invno}.docx", "126", f"{itemname5}")
                                                                                rep_word(f"Invoices//{invno}.docx", "127", f"{item5qty}")
                                                                                rep_word(f"Invoices//{invno}.docx", "128", f"{item5amount}")
                                                                                rep_word(f"Invoices//{invno}.docx", "129", f"{itemname6}")
                                                                                rep_word(f"Invoices//{invno}.docx", "130", f"{item6qty}")
                                                                                rep_word(f"Invoices//{invno}.docx", "131", f"{item6amount}")
                                                                                rep_word(f"Invoices//{invno}.docx", "55", f"{item6qtytotal}")
                                                                                rep_word(f"Invoices//{invno}.docx", "99", f"{item6total}")
                                                                                # print(item2total, "=" , item2qtytotal)
                                                                                if paymode3 == "2":
                                                                                    document = Document(f"Invoices//{invno}.docx")
                                                                                    paragraph = document.add_paragraph()
                                                                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                                                    run = paragraph.add_run()
                                                                                    run.add_picture('qrcode.png', width=Inches(0.8), height=Inches(0.8))

                                                                                    document.save(f"Invoices//{invno}.docx")
                                                                                while True:
                                                                                    if os.path.exists(f"Invoices//{invno}.docx"):
                                                                                        break
                                                                                file = f"Invoices\\{invno}.docx"
                                                                                os.startfile(file)
                                                                                esc = 1
                                                                                break
                                                                            elif add_item7 == 1: # item 7
                                                                                add_item7 = 0
                                                                                while True:
                                                                                    while True:
                                                                                        print("which Product you want to ADD?")
                                                                                        item_code7 = input("Enter Item Code(Example = 120) : ")
                                                                                        if item_code7.isnumeric():
                                                                                            if len(item_code7) == 3:
                                                                                                break
                                                                                            else:
                                                                                                print("Invalid Item Code. Item code Should be 3 Digits")
                                                                                        else:
                                                                                            print("Only numbers are allowed for Item code. Try again!")
                                                                                    while True:
                                                                                        if os.path.exists(f"Items\\{item_code7}.txt"):
                                                                                            break
                                                                                        else:
                                                                                            print("No items in your Inventory with this product code.")
                                                                                            while True:
                                                                                                repet = 0
                                                                                                escape_num = 0
                                                                                                b = input("Do you want to add another product code ? (1 = Y& 0 = No): ")
                                                                                                if b == "1" or b == "0":
                                                                                                    break
                                                                                                else:
                                                                                                    print("Invalid Input! Try Again")
                                                                                                    pass
                                                                                            if b == "1":
                                                                                                repet = 1
                                                                                                break
                                                                                            elif b == "0":
                                                                                                escape_num = 1
                                                                                                break
                                                                                            else:
                                                                                                print("Invalid Input! Try Again")
                                                                                            if escape_num == 1:
                                                                                                break
                                                                                    if escape_num == 1:
                                                                                        escape_num = 0
                                                                                        break
                                                                                    while True:
                                                                                        if repet == 1:
                                                                                            repet = 0
                                                                                            break
                                                                                        f = open(f"Items\\{item_code7}.txt", "r")
                                                                                        itemname7 = f.readline()
                                                                                        itemprice7 = f.readline()
                                                                                        itemstock7 = f.readline()
                                                                                        itemname7 = itemname7.strip()
                                                                                        itemprice7 = itemprice7.strip()
                                                                                        itemstock7 =itemstock7.strip()
                                                                                        itemprice7 = int(itemprice7)
                                                                                        itemstock7 = int(itemstock7)
                                                                                        f.close()
                                                                                        print(f"_ _ _ Product Details _ _ _\nItem Name = {itemname7}\nItem MRP/Unit = {itemprice7}\nItem Stock = {itemstock7}")
                                                                                        while True:
                                                                                            item7qty = input("Enter Quantity : ")
                                                                                            if item7qty.isnumeric():
                                                                                                if int(item7qty) >= 0 and int(item7qty) <100:
                                                                                                    item7qty = int(item7qty)
                                                                                                    if item7qty <= itemstock7:
                                                                                                        new_item_stock = (itemstock7 - item7qty)
                                                                                                        f = open(f"Items\\{item_code7}.txt", "w")
                                                                                                        f.write(f"{itemname7}\n{itemprice7}\n{new_item_stock}")
                                                                                                        f.close()
                                                                                                        print("New Item Stock = ", new_item_stock)
                                                                                                        break
                                                                                                    else:
                                                                                                        print("You don't have enough stock to fulfill the order")
                                                                                                        while True:
                                                                                                            c = input("Do you want to continue with this stock ? (1 = Y& 0 = No): ")
                                                                                                            if c == "1" or c == "0":
                                                                                                                break
                                                                                                            else:
                                                                                                                print("Invalid Input! Try Again")
                                                                                                                pass
                                                                                                        if c == "1":
                                                                                                            pass
                                                                                                        elif c == "0":
                                                                                                            acer = 1
                                                                                                            break
                                                                                                        else:
                                                                                                            print("Invalid Input! Try Again")
                                                                                                else:
                                                                                                    print("Invalid Item Code. Item Quantity Should be less than 2 Digits")
                                                                                            else:
                                                                                                print("Only numbers are allowed for Item code. Try again!")
                                                                                        if acer == 1:
                                                                                            acer = 0
                                                                                            break

                                                                                        print("Item 7 Added")
                                                                                        items = 7
                                                                                        item1amount = item1qty * itemprice1
                                                                                        item2amount = item2qty * itemprice2
                                                                                        item3amount = item3qty * itemprice3
                                                                                        item4amount = item4qty * itemprice4
                                                                                        item5amount = item5qty * itemprice5
                                                                                        item6amount = item6qty * itemprice6
                                                                                        item7amount = item7qty * itemprice7
                                                                                        item7total = item1amount+item2amount+item3amount+item4amount+item5amount+item6amount+item7amount
                                                                                        item7qtytotal = item1qty + item2qty + item3qty + item4qty + item5qty + item6qty + item7qty

                                                                                        print(f"Current Bill Amount is = {item7total}")
                                                                                        while True:
                                                                                            in3 = input("1. Add another item\n2. Generate Incoice\n Enter your Choice : ")
                                                                                            if in3 == "1" or in3 == "2":
                                                                                                if in3 == "1":
                                                                                                    add_item8 = 1
                                                                                                    break
                                                                                                elif in3 == "2":
                                                                                                    generateinvoice7 = 1
                                                                                                    break
                                                                                                else:
                                                                                                    print("Invalid Input! Please Try Again.")
                                                                                            else:
                                                                                                print("Invalid Input! Please Try Again.")

                                                                                        if generateinvoice7 == 1:
                                                                                            generateinvoice7 = 0
                                                                                            while True:
                                                                                                paymode3 = input("Payment Mode:\n1. Cash Payment\n2. Upi Payment\n Enter your Choice : ")
                                                                                                if paymode3 == "1" or paymode3 == "2":
                                                                                                    if paymode3 == "1":
                                                                                                        file2 = "Templetes//7.docx"
                                                                                                        break
                                                                                                    elif paymode3 == "2":
                                                                                                        file2 = "Templetes//7QR.docx"
                                                                                                        link = f"upi://pay?pa=rajlakshmi.mohanty@ybl&pn=abcd&am={item7total}&tn=Tasty Confectionary&cu=INR"
                                                                                                        img = qrcode.make(link)
                                                                                                        img.save("qrcode.png")
                                                                                                        break
                                                                                                else:
                                                                                                    print("Invalid Input! Please Try Again.")
                                                                                            files = os.listdir(directory_path)
                                                                                            num_files = len(files)
                                                                                            invno = num_files + 1
                                                                                            print("Invoice Number is = ", invno)
                                                                                            now = datetime.datetime.now()
                                                                                            current_datetime = now.strftime("%d-%m-%Y        Time - %H:%M:%S")
                                                                                            rep_word(f"{file2}", "111", f"INV {invno}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "112", f"{current_datetime}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "114", f"{itemname1}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "115", f"{item1qty}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "116", f"{item1amount}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "117", f"{itemname2}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "118", f"{item2qty}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "119", f"{item2amount}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "120", f"{itemname3}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "121", f"{item3qty}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "122", f"{item3amount}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "123", f"{itemname4}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "124", f"{item4qty}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "125", f"{item4amount}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "126", f"{itemname5}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "127", f"{item5qty}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "128", f"{item5amount}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "129", f"{itemname6}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "130", f"{item6qty}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "131", f"{item6amount}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "132", f"{itemname7}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "133", f"{item7qty}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "134", f"{item7amount}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "55", f"{item7qtytotal}")
                                                                                            rep_word(f"Invoices//{invno}.docx", "99", f"{item7total}")
                                                                                            # print(item2total, "=" , item2qtytotal)
                                                                                            if paymode3 == "2":
                                                                                                document = Document(f"Invoices//{invno}.docx")
                                                                                                paragraph = document.add_paragraph()
                                                                                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                                                                run = paragraph.add_run()
                                                                                                run.add_picture('qrcode.png', width=Inches(0.8), height=Inches(0.8))

                                                                                                document.save(f"Invoices//{invno}.docx")
                                                                                            while True:
                                                                                                if os.path.exists(f"Invoices//{invno}.docx"):
                                                                                                    break
                                                                                            file = f"Invoices\\{invno}.docx"
                                                                                            os.startfile(file)
                                                                                            esc = 1
                                                                                            break
                                                                                        elif add_item8 == 1: # item 8
                                                                                            add_item8 = 0
                                                                                            while True:
                                                                                                while True:
                                                                                                    print("which Product you want to ADD?")
                                                                                                    item_code8 = input("Enter Item Code(Example = 120) : ")
                                                                                                    if item_code8.isnumeric():
                                                                                                        if len(item_code8) == 3:
                                                                                                            break
                                                                                                        else:
                                                                                                            print("Invalid Item Code. Item code Should be 3 Digits")
                                                                                                    else:
                                                                                                        print("Only numbers are allowed for Item code. Try again!")
                                                                                                while True:
                                                                                                    if os.path.exists(f"Items\\{item_code8}.txt"):
                                                                                                        break
                                                                                                    else:
                                                                                                        print("No items in your Inventory with this product code.")
                                                                                                        while True:
                                                                                                            repet = 0
                                                                                                            escape_num = 0
                                                                                                            b = input("Do you want to add another product code ? (1 = Y& 0 = No): ")
                                                                                                            if b == "1" or b == "0":
                                                                                                                break
                                                                                                            else:
                                                                                                                print("Invalid Input! Try Again")
                                                                                                                pass
                                                                                                        if b == "1":
                                                                                                            repet = 1
                                                                                                            break
                                                                                                        elif b == "0":
                                                                                                            escape_num = 1
                                                                                                            break
                                                                                                        else:
                                                                                                            print("Invalid Input! Try Again")
                                                                                                        if escape_num == 1:
                                                                                                            break
                                                                                                if escape_num == 1:
                                                                                                    escape_num = 0
                                                                                                    break
                                                                                                while True:
                                                                                                    if repet == 1:
                                                                                                        repet = 0
                                                                                                        break
                                                                                                    f = open(f"Items\\{item_code8}.txt", "r")
                                                                                                    itemname8 = f.readline()
                                                                                                    itemprice8 = f.readline()
                                                                                                    itemstock8 = f.readline()
                                                                                                    itemname8 = itemname8.strip()
                                                                                                    itemprice8 = itemprice8.strip()
                                                                                                    itemstock8 =itemstock8.strip()
                                                                                                    itemprice8 = int(itemprice8)
                                                                                                    itemstock8 = int(itemstock8)
                                                                                                    f.close()
                                                                                                    print(f"_ _ _ Product Details _ _ _\nItem Name = {itemname8}\nItem MRP/Unit = {itemprice8}\nItem Stock = {itemstock8}")
                                                                                                    while True:
                                                                                                        item8qty = input("Enter Quantity : ")
                                                                                                        if item8qty.isnumeric():
                                                                                                            if int(item8qty) >= 0 and int(item8qty) <100:
                                                                                                                item8qty = int(item8qty)
                                                                                                                if item8qty <= itemstock8:
                                                                                                                    new_item_stock = (itemstock8 - item8qty)
                                                                                                                    f = open(f"Items\\{item_code8}.txt", "w")
                                                                                                                    f.write(f"{itemname8}\n{itemprice8}\n{new_item_stock}")
                                                                                                                    f.close()
                                                                                                                    print("New Item Stock = ", new_item_stock)
                                                                                                                    break
                                                                                                                else:
                                                                                                                    print("You don't have enough stock to fulfill the order")
                                                                                                                    while True:
                                                                                                                        c = input("Do you want to continue with this stock ? (1 = Y& 0 = No): ")
                                                                                                                        if c == "1" or c == "0":
                                                                                                                            break
                                                                                                                        else:
                                                                                                                            print("Invalid Input! Try Again")
                                                                                                                            pass
                                                                                                                    if c == "1":
                                                                                                                        pass
                                                                                                                    elif c == "0":
                                                                                                                        acer = 1
                                                                                                                        break
                                                                                                                    else:
                                                                                                                        print("Invalid Input! Try Again")
                                                                                                            else:
                                                                                                                print("Invalid Item Code. Item Quantity Should be less than 2 Digits")
                                                                                                        else:
                                                                                                            print("Only numbers are allowed for Item code. Try again!")
                                                                                                    if acer == 1:
                                                                                                        acer = 0
                                                                                                        break

                                                                                                    print("Item 8 Added")
                                                                                                    items = 8
                                                                                                    item1amount = item1qty * itemprice1
                                                                                                    item2amount = item2qty * itemprice2
                                                                                                    item3amount = item3qty * itemprice3
                                                                                                    item4amount = item4qty * itemprice4
                                                                                                    item5amount = item5qty * itemprice5
                                                                                                    item6amount = item6qty * itemprice6
                                                                                                    item7amount = item7qty * itemprice7
                                                                                                    item8amount = item8qty * itemprice8
                                                                                                    item8total = item1amount+item2amount+item3amount+item4amount+item5amount+item6amount+item7amount+item8amount
                                                                                                    item8qtytotal = item1qty + item2qty + item3qty + item4qty + item5qty + item6qty + item7qty + item8qty

                                                                                                    print(f"Current Bill Amount is = {item8total}")
                                                                                                    while True:
                                                                                                        in3 = input("2. Generate Incoice\n Enter your Choice : ")
                                                                                                        if in3 == "1" or in3 == "2":
                                                                                                            if in3 == "1":
                                                                                                                print("Sorry! You can add 8 Items only.")
                                                                                                                # add_item9 = 1
                                                                                                                # break
                                                                                                            elif in3 == "2":
                                                                                                                generateinvoice8 = 1
                                                                                                                break
                                                                                                            else:
                                                                                                                print("Invalid Input! Please Try Again.")
                                                                                                        else:
                                                                                                            print("Invalid Input! Please Try Again.")

                                                                                                    if generateinvoice8 == 1:
                                                                                                        generateinvoice8 = 0
                                                                                                        while True:
                                                                                                            paymode3 = input("Payment Mode:\n1. Cash Payment\n2. Upi Payment\n Enter your Choice : ")
                                                                                                            if paymode3 == "1" or paymode3 == "2":
                                                                                                                if paymode3 == "1":
                                                                                                                    file2 = "Templetes//8.docx"
                                                                                                                    break
                                                                                                                elif paymode3 == "2":
                                                                                                                    file2 = "Templetes//8QR.docx"
                                                                                                                    link = f"upi://pay?pa=rajlakshmi.mohanty@ybl&pn=abcd&am={item8total}&tn=Tasty Confectionary&cu=INR"
                                                                                                                    img = qrcode.make(link)
                                                                                                                    img.save("qrcode.png")
                                                                                                                    break
                                                                                                            else:
                                                                                                                print("Invalid Input! Please Try Again.")
                                                                                                        files = os.listdir(directory_path)
                                                                                                        num_files = len(files)
                                                                                                        invno = num_files + 1
                                                                                                        print("Invoice Number is = ", invno)
                                                                                                        now = datetime.datetime.now()
                                                                                                        current_datetime = now.strftime("%d-%m-%Y        Time - %H:%M:%S")
                                                                                                        rep_word(f"{file2}", "111", f"INV {invno}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "112", f"{current_datetime}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "114", f"{itemname1}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "115", f"{item1qty}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "199", f"{item1amount}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "117", f"{itemname2}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "118", f"{item2qty}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "290", f"{item2amount}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "120", f"{itemname3}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "121", f"{item3qty}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "122", f"{item3amount}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "123", f"{itemname4}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "124", f"{item4qty}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "125", f"{item4amount}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "126", f"{itemname5}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "127", f"{item5qty}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "128", f"{item5amount}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "193", f"{itemname6}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "130", f"{item6qty}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "131", f"{item6amount}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "132", f"{itemname7}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "133", f"{item7qty}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "134", f"{item7amount}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "135", f"{itemname8}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "136", f"{item8qty}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "137", f"{item8amount}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "55", f"{item8qtytotal}")
                                                                                                        rep_word(f"Invoices//{invno}.docx", "99", f"{item8total}")
                                                                                                        # print(item2total, "=" , item2qtytotal)
                                                                                                        if paymode3 == "2":
                                                                                                            document = Document(f"Invoices//{invno}.docx")
                                                                                                            paragraph = document.add_paragraph()
                                                                                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                                                                            run = paragraph.add_run()
                                                                                                            run.add_picture('qrcode.png', width=Inches(0.8), height=Inches(0.8))

                                                                                                            document.save(f"Invoices//{invno}.docx")
                                                                                                        while True:
                                                                                                            if os.path.exists(f"Invoices//{invno}.docx"):
                                                                                                                break
                                                                                                        file = f"Invoices\\{invno}.docx"
                                                                                                        os.startfile(file)
                                                                                                        esc = 1
                                                                                                        break
                                                                                                    if esc == 1:
                                                                                                        break
                                                                                                if esc == 1:
                                                                                                    break
                                                                                            if esc == 1:
                                                                                                break
                                                                                        if esc == 1:
                                                                                            break
                                                                                    if esc == 1:
                                                                                        break
                                                                                if esc == 1:
                                                                                    break
                                                                            if esc == 1:
                                                                                break
                                                                        if esc == 1:
                                                                            break
                                                                    if esc == 1:
                                                                        break
                                                                if esc == 1:
                                                                    break
                                                            if esc == 1:
                                                                break
                                                        if esc == 1:
                                                            break
                                                    if esc == 1:
                                                        break
                                                if esc == 1:
                                                    break
                                            if esc == 1:
                                                break
                                        if esc == 1:
                                            break
                                    if esc == 1:
                                        break
                                if esc == 1:
                                    break
                            if esc == 1:
                                break
                        if esc == 1:
                            break
                    if esc == 1:
                        break
                if esc == 1:
                    break
            if esc == 1:
                break
                    
            # Reprint last Invoice
    elif choice == "2":
        files = os.listdir(directory_path)
        num_files = len(files)
        invno = num_files
        file = f"{directory_path}\\{invno}.docx"
        os.startfile(file, "print")
            # Reprint Invoice
    elif choice == "3":
        while True:
            print("---------------------------")
            print("      Invoice Reprint     ")
            print("---------------------------")
            while True:
                bill = input("Press 0000 To exit \nEnter Bill No. :")
                if bill.isnumeric():
                    if len(bill) > 0 or bill == "0000":
                        if bill == "0000":
                            qwe = 1
                            break
                        else:
                            if os.path.exists(f"Invoices\\{bill}.docx"):
                                file = f"Invoices\\{bill}.docx"
                                os.startfile(file,'print')
                                break
                            else:
                                print("No Invoice in your Invoices with this Invoice No.")
                    else:
                        print("Invalid Item Code. Item code Should be 3 Digits")
                else:
                    print("Only numbers are allowed for Invoice Number. Try again!")
            if qwe == 1:
                qwe = 0
                break
            # View Bills
    elif choice == "4":
        while True:
            print("---------------------------")
            print("        View Invoice       ")
            print("---------------------------")
            while True:
                bill = input("Press 0000 To exit \nEnter Bill No. :")
                if bill.isnumeric():
                    if len(bill) > 0 or bill == "0000":
                        if bill == "0000":
                            qwe = 1
                            break
                        else:
                            if os.path.exists(f"Invoices\\{bill}.docx"):
                                file = f"Invoices\\{bill}.docx"
                                os.startfile(file)
                                break
                            else:
                                print("No Invoice in your Invoices with this Invoice No.")
                    else:
                        print("Invalid Item Code. Item code Should be 3 Digits")
                else:
                    print("Only numbers are allowed for Invoice Number. Try again!")
            if qwe == 1:
                qwe = 0
                break
            # Products Available 
    elif choice == "5":
        print("---------------------------------")
        print("        Available Products       ")
        print("---------------------------------")
        directory_path = "Items\\"
        print("Code    Item Name   Price/Unit  Stock")
        print("-------------------------------------")
        # Iterate through each file in the directory
        for file_name in os.listdir(directory_path):
            if file_name.endswith(".txt"):  # Check if the file is a text file
                file_path = os.path.join(directory_path, file_name)
                with open(file_path, 'r') as file:
                    first_line = file.readline().strip()
                    second_line = file.readline().strip()
                    third_line = file.readline().strip()

                    if int(third_line) == 0:
                        continue
                    file_name = file_name[0:3]
                    length = len(first_line)
                    first_line = first_line + (17-length)*" "
                    print(f"{file_name} : {first_line}: {second_line} \t{third_line}")
                    # Search Products
    elif choice == "6":
        print("---------------------------------")
        print("        Available Products       ")
        print("---------------------------------")
        while True:
            print("Press 0000 to exit")
            while True:
                print("which Product you want to Check?")
                item_code = input("Enter Item Code(Example = 120) : ")
                if item_code == "0000":
                    qwe = 1
                    break

                if item_code.isnumeric():
                    if len(item_code) == 3:
                        if os.path.exists(f"Items\\{item_code}.txt"):
                            f = open(f"Items\\{item_code}.txt", "r")
                            itemname = f.readline()
                            itemprice = f.readline()
                            itemstock = f.readline()
                            itemname = itemname.strip()
                            itemprice = itemprice.strip()
                            itemstock =itemstock.strip()
                            itemprice = int(itemprice)
                            itemstock = int(itemstock)
                            f.close()
                            if itemstock > 0:
                                print("\nThis Item is Available\n")
                            else:
                                print("\nThis Item is not Available\n")
                            print(f"_ _ _ Product Details _ _ _\nItem Name = {itemname}\nItem Price/unit = {itemprice}\nItem Stock = {itemstock}\n")
                    else:
                        print("Invalid Item Code. Item code Should be 3 Digits")
                else:
                    print("Only numbers are allowed for Item code. Try again!")
            if qwe == 1:
                qwe = 0
                break

            # Admin Panel
    elif choice == "7":
        print("---------------------------------")
        print("           Admin Panel           ")
        print("---------------------------------")
        while True:
            admin_code = "5"
            a = input("Enter Admin Pass Code: ")
            if a == admin_code:
                print("Login Success")
                break
            else:
                print("Invalid Input! Try Again")
        while True:
            while True:
                print("1. Add Item")
                print("2. Edit Item")
                print("3. Update Stock")
                print("4. Remove Item")
                print("5. Exit")
                choice = input("Enter Your Choice: ")
                if choice == "1" or choice == "2" or choice == "3" or choice == "4" or choice == "5":
                    break
                else:
                    print("Invalid Input! Try Again")
            if choice == "1":
                # add item
                while True:
                    while True:
                        new_item_code = input("Enter Item Code(Example = 120) : ")
                        if new_item_code.isnumeric():
                            if len(new_item_code) == 3:
                                break
                            else:
                                print("Invalid Item Code. Item code Should be 3 Digits")
                        else:
                            print("Only numbers are allowed for Item code. Try again!")
                    if os.path.exists(f"Items\\{new_item_code}.txt"):
                        print("A item with the same code is available in your inventory.")
                        while True:
                            b = input("Do you want to add another product code ? (1 = Yes & 0 = No): ")
                            if b == "1" or b == "0":
                                if b == "1":
                                    break
                                elif b == "0":
                                    escape_num = 1
                                    break
                                else:
                                    print("Invalid Input! Try Again")
                            else:
                                print("Invalid Input! Try Again")
                                pass
                    else:
                        print("This is a new item. Let's Continue..")
                        break
                    if escape_num == 1:

                        break
                if escape_num == 1:
                    escape_num = 0
                    break
                while True:
                    new_item_name = input("Enter Product Name : ")
                    if len(new_item_name)< 15:
                        break
                    else:
                        print("Only 15 characters are allowed in Item Name!")
                        print("Try Again")
                        pass
                while True:
                    new_item_price = input("Enter MRP/Unit = ")
                    if new_item_price.isnumeric():
                        new_item_price = int(new_item_price)
                        if new_item_price > 0 and new_item_price <= 10000:
                            break
                        else:
                            print("Invalid Input! Try Again")
                    else:
                        print("Invalid Input! Try Again")
                        pass
                while True:
                    new_item_stock = input("Enter opening stock = ")
                    if new_item_stock.isnumeric():
                        new_item_stock = int(new_item_stock)
                        if new_item_stock >= 0 and new_item_stock <= 1000:
                            break
                        else:
                            print("Invalid Input! Try Again")
                    else:
                        print("Invalid Input! Try Again")
                        pass
                new_item_name = new_item_name.capitalize()
                print(f"_ _ _ _ _ Preview _ _ _ _ _\nItem Code = {new_item_code}\nItem Name = {new_item_name}\nItem MRP/Unit = {new_item_price}\nOpening Stock = {new_item_stock}")
                while True:
                    c = input("Do you want to save It? (1 = Yes & 0 = No): ")
                    if c == "1" or c == "0":
                        if c == "1":
                            new_item_name = new_item_name.capitalize()
                            f = open(f"Items\\{new_item_code}.txt", "w")
                            f.write(f"{new_item_name}\n{new_item_price}\n{new_item_stock}")
                            f.close()
                            print("Item registered Successfully")
                            break
                        elif c == "0":
                            print("Item Discarded")
                            break
                        else:
                            print("Invalid Input! Try Again")
                            pass
                    else:
                        print("Invalid Input! Try Again")
             
            elif choice == "2":
                while True:
                    while True:
                        print("which Product you want to Edit?")
                        item_code = input("Enter Item Code(Example = 120) : ")
                        if item_code.isnumeric():
                            if len(item_code) == 3:
                                break
                            else:
                                print("Invalid Item Code. Item code Should be 3 Digits")
                        else:
                            print("Only numbers are allowed for Item code. Try again!")
                    while True:
                        if os.path.exists(f"Items\\{item_code}.txt"):
                            break
                        else:
                            print("No items in your Inventory with this product code.")
                            while True:
                                b = input("Do you want to edit another product code ? (1 = Yes & 0 = No): ")
                                if b == "1" or b == "0":
                                    break
                                else:
                                    print("Invalid Input! Try Again")
                                    pass
                            if b == "1":
                                repet = 1
                                # print(f"repet = {repet}")
                                break
                            elif b == "0":
                                escape_num = 1
                                break
                            else:
                                print("Invalid Input! Try Again")
                            if escape_num == 1:
                                break
                    if escape_num == 1:
                        escape_num = 0
                        break
                    while True:

                        if repet == 1:
                            repet = 0
                            break
                        f = open(f"Items\\{item_code}.txt", "r")
                        old_name = f.readline()
                        old_price = f.readline()
                        old_stock = f.readline()
                        old_name = old_name.strip()
                        old_price = old_price.strip()
                        old_stock = old_stock.strip()
                        print(f"_ _ _ Product Details _ _ _\nItem Name = {old_name}\nItem MRP/Unit = {old_price}")
                        while True:
                            print("Which you want to Edit?\n1. Item Name\n2. MRP/Unit\n3. Both")
                            d = input("Enter your Choice : ")
                            if d == "1" or d == "2" or d == "3":
                                break
                            else:
                                print("Invalid Input! Try Again")
                        if d == "1":
                            while True:
                                new_name = input("Enter New Name : ")
                                if len(new_name) < 15:
                                    break
                                else:
                                    print("Invalid Input! Try Again")
                            new_name = new_name.capitalize()
                            print(f"_ _ _ Preview _ _ _\nItem Name = {new_name}\n Item MRP/Unit = {old_price}")
                            while True:
                                e = input("Do you want to save It? (1 = Yes & 0 = No): ")
                                if e == "1" or e == "0":
                                    if e == "1":
                                        f = open(f"Items\\{item_code}.txt", "w")
                                        f.write(f"{new_name}\n{old_price}\n{old_stock}")
                                        f.close()
                                        print("Item Updated Successfully")
                                        discard = 1
                                        break
                                    elif e == "0":
                                        discard = 1
                                        print("Item Discarded")
                                        break
                                    else:
                                        print("Invalid Input! Try Again")
                                        pass
                                else:
                                    print("Invalid Input! Try Again")
                            if discard == 1:
                                discard = 0
                                while True:
                                    esca = input("Do you want to edit another Item? (1 = Yes & 0 = No): ")
                                    if esca == "1" or esca == "0":
                                        break
                                    else:
                                        print("Invalid Input! Try Again")
                                if esca == "0":
                                    esc = 1
                                break
                        elif d == "2":
                            while True:
                                new_price = input("Enter MRP/Unit = ")
                                if new_price.isnumeric():
                                    new_price = int(new_price)
                                    if new_price > 0 and new_price < 10000:
                                        break
                                    else:
                                        print("Invalid Input! Try Again")
                                else:
                                    print("Invalid Input! Try Again")
                                    pass
                            print(f"_ _ _ Preview _ _ _\nItem Name = {old_name}\n Item MRP/Unit = {new_price}")
                            while True:
                                e = input("Do you want to save It? (1 = Yes & 0 = No): ")
                                if e == "1" or e == "0":
                                    if e == "1":
                                        f = open(f"Items\\{item_code}.txt", "w")
                                        f.write(f"{old_name}\n{new_price}\n{old_stock}")
                                        f.close()
                                        print("Item Updated Successfully")
                                        discard = 1
                                        break
                                    elif e == "0":
                                        discard = 1
                                        print("Item Discarded")
                                        break
                                    else:
                                        print("Invalid Input! Try Again")
                                        pass
                                else:
                                    print("Invalid Input! Try Again")
                            if discard == 1:
                                discard = 0
                                while True:
                                    esca = input("Do you want to edit another Item? (1 = Yes & 0 = No): ")
                                    if esca == "1" or esca == "0":
                                        break
                                    else:
                                        print("Invalid Input! Try Again")
                                if esca == "0":
                                    esc = 1
                                break
                        elif d == "3":
                            while True:
                                new_name = input("Enter New Name : ")
                                if len(new_name) < 20:
                                    break
                                else:
                                    print("Invalid Input! Try Again")
                            while True:
                                new_price = input("Enter MRP/Unit = ")
                                if new_price.isnumeric():
                                    new_price = int(new_price)
                                    if new_price > 0 and new_price < 10000:
                                        break
                                    else:
                                        print("Invalid Input! Try Again")
                                else:
                                    print("Invalid Input! Try Again")
                                    pass
                            new_name = new_name.capitalize()
                            print(f"_ _ _ Preview _ _ _\nItem Name = {new_name}\n Item MRP/Unit = {new_price}")
                            while True:
                                e = input("Do you want to save It? (1 = Yes & 0 = No): ")
                                if e == "1" or e == "0":
                                    if e == "1":
                                        f = open(f"Items\\{item_code}.txt", "w")
                                        f.write(f"{new_name}\n{new_price}\n{old_stock}")
                                        f.close()
                                        print("Item Updated Successfully")
                                        discard = 1
                                        break
                                    elif e == "0":
                                        discard = 1
                                        print("Item Discarded")
                                        break
                                    else:
                                        print("Invalid Input! Try Again")
                                        pass
                                else:
                                    print("Invalid Input! Try Again")
                            if discard == 1:
                                discard = 0
                                while True:
                                    esca = input("Do you want to edit another Item? (1 = Yes & 0 = No): ")
                                    if esca == "1" or esca == "0":
                                        break
                                    else:
                                        print("Invalid Input! Try Again")
                                if esca == "0":
                                    esc = 1
                                break
                    if esc == 1:
                        esc = 0
                        break

            elif choice == "3":
                # update stock
                pass
            elif choice == "4":
                # remove item
                pass
            elif choice == "5":
                # exit
                escape_num = 1
                break
                 
            
            
                    
        

    